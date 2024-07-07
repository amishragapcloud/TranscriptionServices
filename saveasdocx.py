from io import BytesIO
from docx import Document

def save_as_docx(transcript_text):
    doc = Document()
    doc.add_heading('Transcription', level=1)
    doc.add_paragraph(transcript_text)
    byte_io = BytesIO()
    doc.save(byte_io)
    byte_io.seek(0)
    return byte_io